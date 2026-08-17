from __future__ import annotations
from docx import Document
from pathlib import Path
import json, re, sys, datetime

LAYER_META = {
    1: {"name":"POWER", "short":"Power", "description":"External AC/DC power, charger/PSU, battery, charging port and power-button path before motherboard power-rail diagnosis.", "boundary":"If input power reaches the motherboard but the machine remains dead, continue to Layer 2."},
    2: {"name":"MOTHERBOARD / POWER RAILS", "short":"Power Rails", "description":"Motherboard input circuits, standby rails, regulators/VRM, MOSFETs, filtering, sequencing, clocks, reset and PWRGOOD.", "boundary":"When major rails, clock, reset and PWRGOOD are valid but POST still does not begin, continue to Layer 3/4."},
    3: {"name":"CORE HARDWARE", "short":"Core Hardware", "description":"CPU, RAM, GPU and cooling hardware after the board power sequence is reasonably proven.", "boundary":"Hardware faults remain here; firmware initialization/compatibility issues move to Layer 4."},
    4: {"name":"POST / BIOS / UEFI", "short":"POST / Firmware", "description":"Firmware execution, POST, CMOS/RTC, hardware initialization, UEFI boot configuration and handoff toward the bootloader.", "boundary":"Once firmware selects the correct boot path and the OS bootloader is reached, continue to Layer 5."},
    5: {"name":"STORAGE / BOOT", "short":"Storage / Boot", "description":"Storage device/interface, partition table, EFI System Partition, boot files, BCD/bootloader and Windows Boot Manager.", "boundary":"If the OS loader starts and the failure occurs during OS startup, continue to Layer 6."},
    6: {"name":"OS / DRIVER / SOFTWARE", "short":"OS / Software", "description":"Operating system, drivers, device firmware, services, filesystems, applications, security software and system configuration.", "boundary":"This layer begins after power, POST, storage and bootloader stages are reasonably proven."},
}

def slugify(s: str) -> str:
    s = s.lower().replace('—','-').replace('/','-').replace('&','and')
    s = re.sub(r'[^a-z0-9]+','-',s)
    return re.sub(r'-+','-',s).strip('-')

def clean(s: str) -> str:
    s = s.strip()
    s = re.sub(r'[↓↑→─┌┐└┘├┤┬┴│]+',' ',s)
    return re.sub(r'\s+',' ',s).strip()

def parse_code(text: str):
    # Handles "P2/P3", "B1 / B2", "FS1", etc.
    m = re.match(r'^([A-Z]{1,3}\d+(?:\s*/\s*[A-Z]{0,3}\d+)*)\b', text.strip())
    return re.sub(r'\s+','',m.group(1)) if m else None

def expand_codes(code: str):
    if not code: return []
    parts = code.split('/')
    prefix = re.match(r'^[A-Z]+', parts[0]).group(0)
    out=[]
    for p in parts:
        if p.isdigit(): p=prefix+p
        out.append(p)
    return out

def parse_taxonomy(doc: Document, layer_id: int):
    ps=doc.paragraphs
    end=next((i for i,p in enumerate(ps[1:],1) if p.style.name=='Heading 1'), len(ps))
    stack={}
    components=[]
    symptoms=[]
    component_by_path={}
    for p in ps[1:end]:
        t=p.text.rstrip()
        m=re.search(r'([├└]──)\s*(.+)$',t)
        if not m: continue
        depth=m.start()//4
        label=m.group(2).strip()
        code_m=re.match(r'^([A-Z]{1,3}\d+)\.\s*(.+)$',label)
        if code_m:
            code,title=code_m.group(1),code_m.group(2).strip()
            parent_path=[stack[d] for d in sorted(stack) if d < depth]
            component_name=parent_path[-1] if parent_path else LAYER_META[layer_id]['short']
            component_path=' > '.join(parent_path) if parent_path else component_name
            symptoms.append({
                'id': f'L{layer_id}-{code}', 'layerId':layer_id, 'code':code, 'title':title,
                'component':component_name, 'componentPath':component_path,
            })
        else:
            # strip numeric labels such as "1. AC Input / Wall Power"
            name=re.sub(r'^\d+\.\s*','',label).strip()
            stack[depth]=name
            for d in list(stack):
                if d>depth: del stack[d]
            path=' > '.join(stack[d] for d in sorted(stack) if d<=depth)
            key=(layer_id,path)
            if key not in component_by_path:
                cid=f'L{layer_id}-C-{slugify(path)}'
                parent_path=' > '.join(stack[d] for d in sorted(stack) if d<depth)
                parent_id=component_by_path.get((layer_id,parent_path)) if parent_path else None
                comp={'id':cid,'layerId':layer_id,'name':name,'path':path,'depth':depth,'parentId':parent_id}
                components.append(comp); component_by_path[key]=cid
    return components,symptoms

FIELD_MARKERS={
    'symptom': {'symptom','symptoms'},
    'hypothesis': {'form hypothesis','hypothesis'},
    'test': {'perform one test','one test','perform one advanced test'},
    'observe': {'observe','observe result'},
    'verify': {'verify'},
}

def branch_outcomes(raw_flow):
    idx=None; order=None
    for i,line in enumerate(raw_flow):
        norm=clean(line).upper()
        if 'YES' in norm and 'NO' in norm and len(norm) <= 25:
            idx=i
            order=['YES','NO'] if norm.index('YES') < norm.index('NO') else ['NO','YES']
            break
    if idx is None:
        return {'yes':'','no':''}
    cols={order[0]:[],order[1]:[]}
    for line in raw_flow[idx+1:]:
        if re.match(r'^\s*Verify\b',line,re.I): break
        if line.lower().startswith(('reference','references','nguồn')): break
        parts=re.split(r'\s{2,}',line.strip())
        parts=parts[:2]
        if not parts: continue
        values=[clean(x) for x in parts]
        if len(values)==1:
            if values[0]: cols[order[0]].append(values[0])
        else:
            if values[0]: cols[order[0]].append(values[0])
            if values[1]: cols[order[1]].append(values[1])
    def tidy(items):
        # De-duplicate tiny arrow artifacts and preserve source branch wording.
        out=[]
        for x in items:
            if x and (not out or x!=out[-1]): out.append(x)
        return ' → '.join(out)
    return {'yes':tidy(cols['YES']), 'no':tidy(cols['NO'])}

def parse_cases(doc: Document, layer_id: int, taxonomy_symptoms):
    ps=doc.paragraphs
    tax={x['code']:x for x in taxonomy_symptoms}
    cases=[]; current_h1=''
    for i,p in enumerate(ps):
        t=p.text.strip()
        if p.style.name=='Heading 1': current_h1=t
        if p.style.name!='Heading 2': continue
        code=parse_code(t)
        if not code: continue
        j=i+1
        while j<len(ps) and ps[j].style.name not in ('Heading 1','Heading 2'):
            j+=1
        sec=ps[i:j]
        flow_idx=None
        for k,q in enumerate(sec):
            if q.style.name=='Heading 3' and q.text.strip().lower() in ('flow','troubleshooting flow'):
                flow_idx=k; break
        if flow_idx is None: continue
        pre=[]
        for q in sec[1:flow_idx]:
            qt=q.text.strip()
            if qt and q.style.name!='Heading 3' and clean(qt): pre.append(clean(qt))
        raw_flow=[]; notes=[]; refs=[]
        in_after=False
        for q in sec[flow_idx+1:]:
            qt=q.text.strip()
            if not qt: continue
            if q.style.name=='Heading 3':
                in_after=True
            if q.style.name=='Normal (Web)' or in_after or qt.lower().startswith(('reference','references','nguồn')):
                if qt.lower().startswith(('reference','references','nguồn')) or 'http://' in qt or 'https://' in qt:
                    refs.append(qt)
                elif q.style.name=='Normal (Web)':
                    notes.append(qt)
                continue
            raw_flow.append(q.text.rstrip())

        chunks={}; field=None
        branch_started=False
        for line in raw_flow:
            cl=clean(line); low=cl.lower()
            if 'yes' in low and 'no' in low and len(cl)<=25:
                field=None; branch_started=True; continue
            matched=None
            for key,vals in FIELD_MARKERS.items():
                if low in vals: matched=key; break
            if matched:
                field=matched; chunks.setdefault(field,[]); continue
            if field:
                if low.startswith(('narrow down','new hypothesis','find cause','repair','cause identified','move to','continue sequence')) and field!='verify':
                    field=None
                else:
                    chunks[field].append(cl)
        chunks={k:' '.join(v).strip() for k,v in chunks.items()}
        codes=expand_codes(code)
        tax_matches=[tax[c] for c in codes if c in tax]
        first=tax_matches[0] if tax_matches else None
        if '—' in t:
            title=t.split('—',1)[1].strip()
        elif t==code or re.sub(r'\s+','',t)==code:
            title=re.sub(r'^\d+\.\s*','',current_h1).strip()
        else:
            title=first['title'] if first else re.sub(r'^\d+\.\s*','',current_h1).strip()
        branches=branch_outcomes(raw_flow)
        # Extract a useful repair / verification even when branch formatting is parallel.
        repair=''
        for n,line in enumerate(raw_flow):
            cl=clean(line)
            if cl.lower().startswith('repair'):
                if cl.lower()!='repair':
                    repair=cl
                else:
                    bits=[]
                    for nxt in raw_flow[n+1:]:
                        cn=clean(nxt)
                        if not cn: continue
                        if cn.lower().startswith('verify'): break
                        if cn.upper().startswith(('YES','NO')): break
                        bits.append(cn)
                    repair=' '.join(bits[:3])
                break
        verify=chunks.get('verify','')
        if not verify:
            for n,line in enumerate(raw_flow):
                if clean(line).lower()=='verify' and n+1<len(raw_flow):
                    verify=clean(raw_flow[n+1]); break
        cases.append({
            'id': f'L{layer_id}-FLOW-{code.replace("/","-")}',
            'layerId': layer_id,
            'codes': codes,
            'code': code,
            'title': title,
            'component': first['component'] if first else re.sub(r'^\d+\.\s*','',current_h1),
            'componentPath': first['componentPath'] if first else re.sub(r'^\d+\.\s*','',current_h1),
            'sourceSection': current_h1,
            'symptomDescription': ' → '.join(pre),
            'symptom': chunks.get('symptom', first['title'] if first else title),
            'hypothesis': chunks.get('hypothesis',''),
            'test': chunks.get('test',''),
            'observe': chunks.get('observe',''),
            'yesOutcome': branches['yes'],
            'noOutcome': branches['no'],
            'repair': repair,
            'verify': verify,
            'notes': ' '.join(notes).strip(),
            'references': refs,
            'sourceDocument': f'Layer_{layer_id}.docx',
            'status':'verified-source',
            'difficulty': 'Advanced' if any(x in (title+' '+chunks.get('test','')).lower() for x in ['advanced','oscilloscope','programmer','driver verifier','pwrgoo','clock','reset','vgs']) else 'Standard',
            'rawFlow': [clean(x) for x in raw_flow if clean(x)],
        })
    # Add taxonomy symptoms that have no flow section, but do not invent missing diagnostic details.
    covered={c for case in cases for c in case['codes']}
    for s in taxonomy_symptoms:
        if s['code'] not in covered:
            cases.append({
                'id':f"L{layer_id}-FLOW-{s['code']}", 'layerId':layer_id, 'codes':[s['code']], 'code':s['code'],
                'title':s['title'], 'component':s['component'], 'componentPath':s['componentPath'],
                'sourceSection':'Taxonomy only', 'symptomDescription':s['title'], 'symptom':s['title'],
                'hypothesis':'','test':'','observe':'','yesOutcome':'','noOutcome':'','repair':'','verify':'',
                'notes':'The source taxonomy lists this symptom, but the uploaded document does not provide a standalone structured diagnostic flow for it.',
                'references':[], 'sourceDocument':f'Layer_{layer_id}.docx', 'status':'catalog-only', 'difficulty':'Standard','rawFlow':[]
            })
    return cases

def main(paths, out_path):
    layers=[]; components=[]; symptoms=[]; flows=[]
    for layer_id,path in enumerate(paths,1):
        doc=Document(path)
        comps,syms=parse_taxonomy(doc,layer_id)
        layer=dict({'id':layer_id},**LAYER_META[layer_id])
        layer['sourceDocument']=Path(path).name
        layers.append(layer); components.extend(comps); symptoms.extend(syms)
        flows.extend(parse_cases(doc,layer_id,syms))
    # Map taxonomy symptoms to a flow (including grouped cases).
    flow_by_code={}
    for f in flows:
        for c in f['codes']: flow_by_code[(f['layerId'],c)]=f['id']
    for s in symptoms: s['flowId']=flow_by_code.get((s['layerId'],s['code']))
    # Counts
    for layer in layers:
        layer['componentCount']=sum(1 for c in components if c['layerId']==layer['id'])
        layer['symptomCount']=sum(1 for s in symptoms if s['layerId']==layer['id'])
        layer['flowCount']=sum(1 for f in flows if f['layerId']==layer['id'] and f['status']=='verified-source')
    db={
        'meta':{
            'name':'PC Diagnostic Atlas','version':'0.1.0',
            'generatedAt':datetime.datetime.now(datetime.timezone.utc).isoformat(),
            'sourceDocuments':[Path(x).name for x in paths],
            'method':'Taxonomy and diagnostic flows extracted from six uploaded DOCX source documents. Missing standalone flows are marked catalog-only rather than invented.'
        },
        'layers':layers,'components':components,'symptoms':symptoms,'flows':flows,'sessions':[]
    }
    Path(out_path).write_text(json.dumps(db,ensure_ascii=False,indent=2),encoding='utf-8')
    print(f"Wrote {out_path}: {len(layers)} layers, {len(components)} components, {len(symptoms)} taxonomy symptoms, {len(flows)} flows")
    print('Verified source flows:',sum(f['status']=='verified-source' for f in flows),'catalog-only:',sum(f['status']=='catalog-only' for f in flows))

if __name__=='__main__':
    if len(sys.argv)<8:
        raise SystemExit('Usage: build_seed_from_docx.py Layer_1.docx ... Layer_6.docx output.json')
    main(sys.argv[1:7],sys.argv[7])
